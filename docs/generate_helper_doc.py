"""
Generates Helper Trip Commission Module documentation in .docx and .pdf formats.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, KeepTogether, ListFlowable, ListItem
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY

import os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUT_DIR, "Helper_Trip_Commission_Module.docx")
PDF_PATH = os.path.join(OUT_DIR, "Helper_Trip_Commission_Module.pdf")


# ==================== WORD DOCUMENT ====================
def build_docx():
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x5C)
        return h

    def add_para(text, bold=False, size=11, color=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        return p

    def add_bullet(text):
        p = doc.add_paragraph(text, style="List Bullet")
        return p

    def add_table(headers, rows, col_widths=None):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for p in hdr_cells[i].paragraphs:
                for r in p.runs:
                    r.bold = True
        for r_idx, row in enumerate(rows, start=1):
            for c_idx, val in enumerate(row):
                table.rows[r_idx].cells[c_idx].text = str(val)
        if col_widths:
            for row in table.rows:
                for idx, width in enumerate(col_widths):
                    row.cells[idx].width = width
        return table

    # ----- TITLE PAGE -----
    title = doc.add_heading("Helper Trip Commission Module", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x5C)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Design & Flow Document")
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(0x5B, 0x6B, 0x8C)

    doc.add_paragraph()

    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.style = "Light List Accent 1"
    meta_data = [
        ("Product", "WINIT CommissionIQ"),
        ("Module", "Section 6.3 — Helper Trip Commission Engine"),
        ("Live URL", "https://commission-app-six.vercel.app"),
        ("Prepared by", "Engineering Team"),
        ("Version / Date", "v1.0  —  11 April 2026"),
    ]
    for i, (k, v) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = k
        meta_table.rows[i].cells[1].text = v
        for run in meta_table.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True

    doc.add_page_break()

    # ----- 1. EXECUTIVE SUMMARY -----
    add_heading("1. Executive Summary", 1)
    doc.add_paragraph(
        "The Helper Trip Commission module rewards field helpers (delivery drivers, loaders, "
        "unloaders, van sales crew) on a per-trip basis, where the per-person rate depends on "
        "how many helpers shared the trip and how many days the trip lasted."
    )
    doc.add_paragraph(
        "Business rule: Fewer helpers on a trip → higher per-person rate. More helpers → "
        "lower per-person rate (shared reward). Multi-day trips multiply the rate by the "
        "number of days worked."
    )
    doc.add_paragraph(
        "This module is fully integrated into the existing 13-step commission calculation "
        "pipeline and plugs into the standard approval workflow, audit trail, and payout process."
    )

    # ----- 2. BUSINESS REQUIREMENT -----
    add_heading("2. Business Requirement", 1)
    add_table(
        ["Concern", "Requirement"],
        [
            ["Incentive goal", "Pay helpers based on actual field work (trips completed), not just attendance"],
            ["Solo vs team", "Solo workers carry the full load, so they earn more per trip than team members who share the work"],
            ["Multi-day trips", "Long-distance trips (2+ days) should pay proportionally"],
            ["Configurability", "Rates must be editable per plan; no code changes required"],
            ["Auditability", "Every paid trip must be traceable — who, when, what team size, what rate"],
            ["Integration", "Must flow through existing caps, multipliers, penalties, and approvals"],
        ],
    )

    # ----- 3. COMMISSION FORMULA -----
    add_heading("3. Commission Formula", 1)
    formula = doc.add_paragraph()
    run = formula.add_run(
        "per_person_earned_per_trip = rate_per_person_per_day(team_size) × days_count\n"
        "total_helper_commission    = Σ (per_person_earned_per_trip) for all trips in the period"
    )
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x5C)

    add_heading("Default Rate Table (editable per plan)", 2)
    add_table(
        ["Team Size", "Rate / Person / Day (AED)", "Rationale"],
        [
            ["1 (solo)", "12", "Full rate — no sharing"],
            ["2 (pair)", "7", "Shared workload"],
            ["3 (team)", "5", "Split further"],
            ["4+ (large team)", "4", "Lowest per person"],
        ],
    )
    note = doc.add_paragraph()
    note_run = note.add_run(
        "Per-person rate decreases as team grows — total trip cost to company still rises slightly "
        "(e.g. 1×12 = 12 vs 3×5 = 15), but per-helper earnings go down."
    )
    note_run.italic = True
    note_run.font.size = Pt(10)

    # ----- 4. WORKED EXAMPLES -----
    add_heading("4. Worked Examples", 1)

    add_heading("Example 1 — Solo single-day trip", 2)
    add_bullet("1 helper, 1 day")
    add_bullet("12 AED × 1 day = 12 AED for the helper")

    add_heading("Example 2 — Paired 3-day long-haul", 2)
    add_bullet("2 helpers, 3 days")
    add_bullet("7 AED × 3 days = 21 AED each (2 × 21 = 42 AED total company cost)")

    add_heading("Example 3 — Team of 3, single-day", 2)
    add_bullet("3 helpers, 1 day")
    add_bullet("5 AED × 1 day = 5 AED each (3 × 5 = 15 AED total)")

    add_heading("Example 4 — Khalid Omar's January 2026 (actual live data)", 2)
    add_table(
        ["Trip type", "Count", "Days", "Rate/day", "Subtotal"],
        [
            ["Solo trips", "10", "10", "12", "120 AED"],
            ["Paired trips", "8", "8", "7", "56 AED"],
            ["TOTAL", "18", "18", "", "176 AED"],
        ],
    )

    doc.add_page_break()

    # ----- 5. SYSTEM ARCHITECTURE -----
    add_heading("5. System Architecture", 1)
    doc.add_paragraph("Four-phase flow from setup to payout:")
    p = doc.add_paragraph()
    arch = p.add_run(
        "PLAN SETUP  →  DAILY OPS  →  CALCULATION  →  PAYOUT\n"
        "(Rate Table)    (Trip Log)     (Pipeline)       (Approval)\n\n"
        "Everything routes through the AUDIT TRAIL at each step."
    )
    arch.font.name = "Consolas"
    arch.font.size = Pt(10)

    add_heading("Data Model", 2)
    doc.add_paragraph("Three new tables introduced:")
    add_bullet("helper_trip_rates — Configurable rate tiers (plan-scoped or global default)")
    add_bullet("trips — Trip header with start/end dates, period, status, distance, stops")
    add_bullet("trip_participants — Many-to-many link between trips and employees")

    # ----- 6. END-TO-END USER FLOW -----
    add_heading("6. End-to-End User Flow", 1)

    add_heading("Phase 1 — Plan Setup (one-time per plan)", 2)
    doc.add_paragraph("Actor: Plan Administrator / Compensation Manager")
    doc.add_paragraph("Location: Plans → [Plan Name] → Helper Trips tab")
    for step in [
        "Open an existing commission plan (or create a new one)",
        "Click the Helper Trips tab",
        "In Rate Table by Team Size, review the tiers (default: 1→12, 2→7, 3→5, 4→4 AED/person/day)",
        "Click Apply Defaults to restore if edited",
        "Edit tiers by changing rate values, add new tiers with + Add Tier, remove with trash icon",
        "Click Save Rates to persist",
    ]:
        add_bullet(step)

    add_heading("Phase 2 — Daily Operations (recorded as trips happen)", 2)
    doc.add_paragraph("Actor: Supervisor / Dispatcher / Field Ops")
    doc.add_paragraph("Location: Plans → [Plan Name] → Helper Trips → Trip Log")
    for step in [
        "Click + Log New Trip",
        "Fill in Trip Number, Start Date, End Date (optional for multi-day), Stops, Distance",
        "Click participant chips to select helpers (1, 2, 3, or more)",
        "See live preview: e.g. 2 helpers · 3 days · Rate 7 AED = 21 AED each",
        "Click Create Trip — appears in the Trip Log table",
    ]:
        add_bullet(step)

    doc.add_paragraph("Trip Log filters:")
    add_bullet("Period dropdown — shows months with trips and counts")
    add_bullet("Month picker — choose any month directly")
    add_bullet("All Periods option — shows everything across months")

    add_heading("Phase 3 — Calculation (automated, at period close)", 2)
    doc.add_paragraph("Actor: Commission Administrator")
    doc.add_paragraph("Location: Calculate page")
    for step in [
        "Pick period from header calendar",
        "Select plan from dropdown",
        "Optionally select a specific employee",
        "Click ▶ Run Calculation",
    ]:
        add_bullet(step)

    doc.add_paragraph("The 13-step calculation pipeline runs automatically:")
    pipeline = [
        "1. Fetch transactions (sales, returns, collections)",
        "2. Apply mapping filters (include/exclude rules)",
        "2.5. Eligibility check (min sales, max returns, etc.)",
        "3. KPI achievement (target vs actual)",
        "4. Determine slab (tier lookup)",
        "5. KPI payout (base × slab rate)",
        "6. Apply weight (× weight %)",
        "7. AGGREGATE KPIs + HELPER TRIP BONUS  ← integration point",
        "8. Apply multipliers (growth, strategic)",
        "9. Apply penalties (high returns, audit fail)",
        "10. Apply caps (max per plan, % of salary)",
        "11. Store payout",
        "12. Create approval entry",
        "13. Complete run",
    ]
    for s in pipeline:
        add_bullet(s)

    doc.add_paragraph("Helper integration details (Step 7):")
    for s in [
        "Pipeline queries the trips table joined with trip_participants for the employee and period",
        "For each trip, looks up the correct rate from helper_trip_rates using team size",
        "Multiplies rate × days_count per trip",
        "Sums across all trips",
        "Adds to gross_payout BEFORE multipliers/penalties/caps are applied",
        "Persists helper_trip_bonus + full breakdown in calculation_details JSON",
    ]:
        add_bullet(s)

    add_heading("Phase 4 — Approval Workflow", 2)
    doc.add_paragraph("After calculation, each payout enters the workflow:")
    p = doc.add_paragraph()
    p.add_run("submitted → manager_approved → finance_approved → hr_approved → locked").font.name = "Consolas"
    for s in [
        "Each stage requires a named approver",
        "Rejections require a mandatory reason",
        "Locked payouts cannot be edited",
        "Reopening requires super-admin privileges",
        "Every transition is logged in approval_log (immutable)",
    ]:
        add_bullet(s)

    add_heading("Phase 5 — Audit & Governance", 2)
    doc.add_paragraph("Every change is tracked in the audit_trail table:")
    add_table(
        ["Entity Type", "Tracked Actions"],
        [
            ["trip", "created, updated, deleted"],
            ["helper_rates", "updated (rate changes)"],
            ["employee_payouts", "calculated, approved, locked"],
        ],
    )

    doc.add_page_break()

    # ----- 7. RATE LOOKUP LOGIC -----
    add_heading("7. Rate Lookup Logic", 1)
    doc.add_paragraph("Tiered lookup — finds the highest team_size tier ≤ actual team size.")
    doc.add_paragraph("Behavior:")
    for s in [
        "Team of 1 → picks tier 1 (12 AED)",
        "Team of 2 → picks tier 2 (7 AED)",
        "Team of 3 → picks tier 3 (5 AED)",
        "Team of 5 → picks tier 4 (4 AED, since no tier 5 defined)",
        "Team of 100 → still picks tier 4 (highest defined ≤ 100)",
    ]:
        add_bullet(s)
    doc.add_paragraph(
        "Plan-scoped rates override global defaults. If a plan has its own helper_trip_rates rows, "
        "those are used; otherwise the plan_id IS NULL defaults kick in."
    )

    # ----- 8. API REFERENCE -----
    add_heading("8. API Reference", 1)
    add_table(
        ["Method", "Endpoint", "Purpose"],
        [
            ["GET", "/api/trips", "List trips (filters: employee_id, period, status)"],
            ["GET", "/api/trips/:id", "Single trip with participants"],
            ["POST", "/api/trips", "Create trip (body: dates, participants, stops, distance)"],
            ["PUT", "/api/trips/:id", "Update trip or participants"],
            ["DELETE", "/api/trips/:id", "Remove trip"],
            ["GET", "/api/trips/rates/:planId", "Get rate config (plan-specific or default)"],
            ["PUT", "/api/trips/rates/:planId", "Update rate tiers"],
            ["GET", "/api/trips/commission/preview", "Preview commission for employee+period without posting"],
        ],
    )

    # ----- 9. EDGE CASES -----
    add_heading("9. Edge Cases Handled", 1)
    add_table(
        ["Scenario", "Behavior"],
        [
            ["Trip with no end date", "Treated as 1-day (days_count = 1)"],
            ["End date before start date", "Falls back to 1 day"],
            ["Team size exceeds all tiers", "Uses the highest defined tier"],
            ["No rates configured for plan", "Falls back to global defaults"],
            ["Trip status = cancelled", "Excluded from commission calculation"],
            ["Employee on multiple trips same day", "All counted independently"],
            ["Rate changed mid-period", "Recalculation uses current rate (locked plans preserve historical rates)"],
            ["Empty trip log for period", "Returns total_commission: 0 cleanly"],
        ],
    )

    # ----- 10. INTEGRATION -----
    add_heading("10. Integration with Existing Commission Logic", 1)
    doc.add_paragraph(
        "The helper bonus flows through the standard pipeline, so it inherits all existing behaviors:"
    )
    add_table(
        ["Pipeline Stage", "Effect on Helper Bonus"],
        [
            ["Eligibility Check", "If employee fails min_sales etc., net_payout = 0 — helper bonus zeroed"],
            ["Multipliers", "Bonus multipliers apply to gross (which now includes helper)"],
            ["Penalties", "Penalty % deducts from total including helper"],
            ["Caps", "max_per_plan and percent_of_salary caps apply to final total"],
            ["Splits", "If employee's role is in a split rule, their share applies to final total"],
            ["Approval", "Part of the same approval flow"],
            ["Audit trail", "Full traceability with JSON breakdown in calculation_details"],
        ],
    )
    note = doc.add_paragraph()
    nr = note.add_run(
        "Important: If an employee's total is capped (e.g. 150% of salary), the helper bonus is added "
        "to gross before the cap is enforced. In some cases the cap may consume the helper bonus "
        "entirely — this is intentional and matches how all other commission components behave."
    )
    nr.italic = True

    doc.add_page_break()

    # ----- 11. WORKED EXAMPLE -----
    add_heading("11. Complete Worked Example", 1)
    doc.add_paragraph("Setup: Plan plan-01 'Salesman Monthly Incentive'")
    doc.add_paragraph("Period: 2026-04")
    doc.add_paragraph("Employee: Khalid Omar (emp-003), Van Sales Driver, salary 4,500 AED")

    add_heading("Step 1 — Rates configured", 2)
    add_table(
        ["Team Size", "Rate AED/day"],
        [["1", "12"], ["2", "7"], ["3", "5"], ["4+", "4"]],
    )

    add_heading("Step 2 — Trips logged for April", 2)
    add_table(
        ["Trip #", "Start", "End", "Days", "Team", "Rate", "Earned"],
        [
            ["TRIP-APR-001", "Apr 5", "—", "1", "1 solo", "12", "12"],
            ["TRIP-APR-002", "Apr 6", "Apr 8", "3", "2 pair", "7", "21"],
            ["TRIP-APR-003", "Apr 9", "—", "1", "1 solo", "12", "12"],
            ["TEST-MULTIDAY", "Apr 10", "Apr 12", "3", "1 solo", "12", "36"],
            ["TOTAL", "", "", "8", "", "", "81"],
        ],
    )

    add_heading("Step 3 — Calculation runs", 2)
    doc.add_paragraph("Pipeline execution summary:")
    for s in [
        "Step 1: Fetch 104 transactions + 4 trips = 108 records",
        "Step 2: Apply filters, 104 pass",
        "Step 3: KPI Revenue: 82,341 / 50,000 = 164.68%",
        "Step 4: Tier 5 (120%+) → 12% rate",
        "Step 5: 2000 × 12% = 240",
        "Step 6: 240 × 40% weight = 96",
        "Step 7: AGGREGATE: 48,636 KPI sum + 81 helper trip bonus = 48,717.60 gross",
        "Step 8: No multipliers triggered",
        "Step 9: No penalties triggered",
        "Step 10: Caps: no cap hit → net = 48,717.60",
        "Step 11: Stored",
        "Step 12: Approval submitted",
        "Step 13: Run completed",
    ]:
        add_bullet(s)

    add_heading("Step 4 — UI result", 2)
    add_table(
        ["Metric", "Value"],
        [
            ["Gross Payout", "AED 48,717.60"],
            ["Helper Trip Bonus", "+AED 81 (4 trips · 8 days)"],
            ["Multiplier", "+AED 0"],
            ["Penalty", "-AED 0"],
            ["Net Payout", "AED 48,717.60"],
        ],
    )

    add_heading("Step 5 — Approval", 2)
    doc.add_paragraph(
        "Moves through: submitted → manager_approved → finance_approved → hr_approved → locked → "
        "sent to payroll."
    )

    # ----- 12. ACCEPTANCE CRITERIA -----
    add_heading("12. Acceptance Criteria", 1)
    doc.add_paragraph("All criteria met:")
    for s in [
        "✓ Rate table editable with visual live preview",
        "✓ Solo rate is higher than pair rate (default: 12 vs 7)",
        "✓ Team rates progressively lower (5 → 4)",
        "✓ Multi-day trips multiply correctly (rate × days)",
        "✓ Live form shows calculation before saving a trip",
        "✓ Trip log filtered by period with 'all periods' option",
        "✓ Pipeline Step 7 integrates helper bonus without breaking other steps",
        "✓ Net payout visible in UI with helper bonus highlighted",
        "✓ Approval workflow intact",
        "✓ Audit trail captures helper commission details",
        "✓ Works across multiple plans with plan-scoped overrides",
        "✓ Commission Preview API for pre-calc verification",
    ]:
        add_bullet(s)

    # ----- 13. NON-FUNCTIONAL -----
    add_heading("13. Non-Functional Requirements Satisfied", 1)
    add_table(
        ["NFR", "How"],
        [
            ["Scalability", "Trip queries indexed on employee_id + period; pipeline uses batch per employee"],
            ["Performance", "Rate lookup cached per run; tag context built once per run"],
            ["Auditability", "Full JSON breakdown persisted; every trip has stable UUID"],
            ["Configurability", "All rates editable via UI + API; no code deploy needed"],
            ["Multi-currency", "Rate table has currency column; uses standard exchange_rates"],
            ["Multi-plan", "Rates can be plan-scoped (override default)"],
            ["Reversibility", "Trip deletion supported; historical runs retain snapshot"],
        ],
    )

    # ----- 14. FUTURE ENHANCEMENTS -----
    add_heading("14. Future Enhancements (not in v1)", 1)
    for s in [
        "GPS validation — auto-reject trips where GPS doesn't confirm route completion",
        "Automatic trip creation — pull from telematics / dispatch system via webhook",
        "Approval on trip level — supervisor must approve trip before it counts for commission",
        "Rate variations by route/distance — long-distance trips get bonus multiplier",
        "Seasonal rates — higher rates during peak periods (Ramadan, end-of-year)",
        "Bonus for perfect attendance — additional multiplier for drivers with 100% beat compliance",
    ]:
        add_bullet(s)

    # ----- 15. CONTACT -----
    add_heading("15. Contact & Environment", 1)
    add_table(
        ["Field", "Value"],
        [
            ["Project", "WINIT CommissionIQ Helper Trip Module"],
            ["Environment", "Production (Vercel, Mumbai region)"],
            ["Database", "Turso (libSQL) — commission-lochan.aws-ap-south-1.turso.io"],
            ["Deployment URL", "https://commission-app-six.vercel.app"],
            ["Status", "Live and operational"],
        ],
    )

    doc.add_paragraph()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    er = end.add_run("— End of Document —")
    er.italic = True
    er.font.color.rgb = RGBColor(0x5B, 0x6B, 0x8C)

    doc.save(DOCX_PATH)
    print(f"[OK] DOCX written: {DOCX_PATH}")


# ==================== PDF DOCUMENT ====================
def build_pdf():
    doc = SimpleDocTemplate(
        PDF_PATH,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title="Helper Trip Commission Module",
        author="WINIT CommissionIQ",
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "Title", parent=styles["Title"],
        fontSize=24, textColor=colors.HexColor("#1F2D5C"),
        alignment=TA_CENTER, spaceAfter=6
    )
    subtitle_style = ParagraphStyle(
        "Subtitle", parent=styles["Normal"],
        fontSize=14, textColor=colors.HexColor("#5B6B8C"),
        alignment=TA_CENTER, spaceAfter=20
    )
    h1 = ParagraphStyle(
        "H1", parent=styles["Heading1"],
        fontSize=16, textColor=colors.HexColor("#1F2D5C"),
        spaceBefore=16, spaceAfter=8, fontName="Helvetica-Bold"
    )
    h2 = ParagraphStyle(
        "H2", parent=styles["Heading2"],
        fontSize=13, textColor=colors.HexColor("#344E86"),
        spaceBefore=10, spaceAfter=6, fontName="Helvetica-Bold"
    )
    body = ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontSize=10.5, leading=14, alignment=TA_JUSTIFY, spaceAfter=6
    )
    bullet_style = ParagraphStyle(
        "Bullet", parent=body,
        leftIndent=14, bulletIndent=2, spaceAfter=3
    )
    code_style = ParagraphStyle(
        "Code", parent=body,
        fontName="Courier", fontSize=9.5,
        textColor=colors.HexColor("#1F2D5C"),
        backColor=colors.HexColor("#F5F7FA"),
        borderPadding=6, spaceAfter=10, spaceBefore=4,
    )
    note_style = ParagraphStyle(
        "Note", parent=body,
        fontSize=9.5, textColor=colors.HexColor("#5B6B8C"),
        fontName="Helvetica-Oblique"
    )

    def mk_table(data, col_widths=None, header=True):
        t = Table(data, colWidths=col_widths, hAlign="LEFT")
        style = [
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D5DAE5")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("FONTSIZE", (0, 0), (-1, -1), 9.5),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8F9FB")]),
        ]
        if header:
            style += [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F2D5C")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ]
        t.setStyle(TableStyle(style))
        return t

    story = []

    # TITLE PAGE
    story.append(Spacer(1, 3 * cm))
    story.append(Paragraph("Helper Trip Commission Module", title_style))
    story.append(Paragraph("Design &amp; Flow Document", subtitle_style))
    story.append(Spacer(1, 1 * cm))

    meta = [
        ["Product", "WINIT CommissionIQ"],
        ["Module", "Section 6.3 — Helper Trip Commission Engine"],
        ["Live URL", "https://commission-app-six.vercel.app"],
        ["Prepared by", "Engineering Team"],
        ["Version", "v1.0"],
        ["Date", "11 April 2026"],
    ]
    story.append(mk_table(meta, col_widths=[4.5 * cm, 10 * cm], header=False))
    story.append(PageBreak())

    # 1. EXECUTIVE SUMMARY
    story.append(Paragraph("1. Executive Summary", h1))
    story.append(Paragraph(
        "The Helper Trip Commission module rewards field helpers (delivery drivers, loaders, "
        "unloaders, van sales crew) on a <b>per-trip basis</b>, where the per-person rate depends on "
        "<b>how many helpers shared the trip</b> and <b>how many days the trip lasted</b>.", body))
    story.append(Paragraph(
        "<b>Business rule:</b> Fewer helpers on a trip → higher per-person rate. "
        "More helpers → lower per-person rate (shared reward). "
        "Multi-day trips multiply the rate by the number of days worked.", body))
    story.append(Paragraph(
        "This module is fully integrated into the existing 13-step commission calculation "
        "pipeline and plugs into the standard approval workflow, audit trail, and payout process.", body))

    # 2. BUSINESS REQUIREMENT
    story.append(Paragraph("2. Business Requirement", h1))
    story.append(mk_table([
        ["Concern", "Requirement"],
        ["Incentive goal", "Pay helpers based on actual field work (trips completed), not attendance"],
        ["Solo vs team", "Solo workers earn more per trip; team members share the reward"],
        ["Multi-day trips", "Long-distance trips (2+ days) should pay proportionally"],
        ["Configurability", "Rates must be editable per plan; no code changes required"],
        ["Auditability", "Every paid trip must be traceable — who, when, team size, rate"],
        ["Integration", "Must flow through caps, multipliers, penalties, and approvals"],
    ], col_widths=[3.5 * cm, 12 * cm]))

    # 3. FORMULA
    story.append(Paragraph("3. Commission Formula", h1))
    story.append(Paragraph(
        "per_person_earned_per_trip = rate_per_person_per_day(team_size) × days_count<br/>"
        "total_helper_commission = Σ (per_person_earned_per_trip) for all trips in the period",
        code_style))

    story.append(Paragraph("Default Rate Table (editable per plan)", h2))
    story.append(mk_table([
        ["Team Size", "Rate / Person / Day (AED)", "Rationale"],
        ["1 (solo)", "12", "Full rate — no sharing"],
        ["2 (pair)", "7", "Shared workload"],
        ["3 (team)", "5", "Split further"],
        ["4+ (large team)", "4", "Lowest per person"],
    ], col_widths=[3.5 * cm, 4.5 * cm, 7.5 * cm]))
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        "<i>Per-person rate decreases as team grows — total trip cost to company still rises slightly "
        "(1×12 = 12 vs 3×5 = 15), but per-helper earnings go down.</i>", note_style))

    # 4. WORKED EXAMPLES
    story.append(Paragraph("4. Worked Examples", h1))

    story.append(Paragraph("Example 1 — Solo single-day trip", h2))
    story.append(Paragraph("• 1 helper, 1 day", bullet_style))
    story.append(Paragraph("• <b>12 AED × 1 day = 12 AED</b> for the helper", bullet_style))

    story.append(Paragraph("Example 2 — Paired 3-day long-haul", h2))
    story.append(Paragraph("• 2 helpers, 3 days", bullet_style))
    story.append(Paragraph("• <b>7 AED × 3 days = 21 AED</b> each (2 × 21 = 42 AED total)", bullet_style))

    story.append(Paragraph("Example 3 — Team of 3, single-day", h2))
    story.append(Paragraph("• 3 helpers, 1 day", bullet_style))
    story.append(Paragraph("• <b>5 AED × 1 day = 5 AED</b> each (3 × 5 = 15 AED total)", bullet_style))

    story.append(Paragraph("Example 4 — Khalid Omar's January 2026 (actual live data)", h2))
    story.append(mk_table([
        ["Trip type", "Count", "Days", "Rate/day", "Subtotal"],
        ["Solo trips", "10", "10", "12", "120 AED"],
        ["Paired trips", "8", "8", "7", "56 AED"],
        ["TOTAL", "18", "18", "", "176 AED"],
    ], col_widths=[3.5 * cm, 2 * cm, 2 * cm, 2.5 * cm, 3 * cm]))

    story.append(PageBreak())

    # 5. ARCHITECTURE
    story.append(Paragraph("5. System Architecture", h1))
    story.append(Paragraph("Four-phase flow from setup to payout:", body))
    story.append(Paragraph(
        "PLAN SETUP → DAILY OPS → CALCULATION → PAYOUT<br/>"
        "(Rate Table)  (Trip Log)   (Pipeline)    (Approval)<br/><br/>"
        "Everything routes through the AUDIT TRAIL at each step.", code_style))

    story.append(Paragraph("Data Model", h2))
    story.append(Paragraph("Three new tables introduced:", body))
    for s in [
        "<b>helper_trip_rates</b> — Configurable rate tiers (plan-scoped or global default)",
        "<b>trips</b> — Trip header with start/end dates, period, status, distance, stops",
        "<b>trip_participants</b> — Many-to-many link between trips and employees",
    ]:
        story.append(Paragraph("• " + s, bullet_style))

    # 6. USER FLOW
    story.append(Paragraph("6. End-to-End User Flow", h1))

    story.append(Paragraph("Phase 1 — Plan Setup (one-time per plan)", h2))
    story.append(Paragraph("<b>Actor:</b> Plan Administrator / Compensation Manager", body))
    story.append(Paragraph("<b>Location:</b> Plans → [Plan Name] → Helper Trips tab", body))
    for s in [
        "Open an existing commission plan (or create a new one)",
        "Click the Helper Trips tab",
        "Review rate tiers (default: 1→12, 2→7, 3→5, 4→4 AED)",
        "Edit tiers, add new ones, or click Apply Defaults",
        "Click Save Rates to persist",
    ]:
        story.append(Paragraph("• " + s, bullet_style))

    story.append(Paragraph("Phase 2 — Daily Operations", h2))
    story.append(Paragraph("<b>Actor:</b> Supervisor / Dispatcher / Field Ops", body))
    story.append(Paragraph("<b>Location:</b> Plans → [Plan Name] → Helper Trips → Trip Log", body))
    for s in [
        "Click + Log New Trip",
        "Fill Trip Number, Start Date, End Date (optional), Stops, Distance",
        "Click participant chips to select helpers",
        "See live preview showing calculated commission",
        "Click Create Trip — appears in Trip Log table",
    ]:
        story.append(Paragraph("• " + s, bullet_style))

    story.append(Paragraph("Phase 3 — Calculation (automated, at period close)", h2))
    story.append(Paragraph("<b>Actor:</b> Commission Administrator. <b>Location:</b> Calculate page", body))
    story.append(Paragraph("The 13-step pipeline runs automatically:", body))
    pipeline = [
        "1. Fetch transactions (sales, returns, collections)",
        "2. Apply mapping filters (include/exclude rules)",
        "2.5. Eligibility check (min sales, max returns, etc.)",
        "3. KPI achievement (target vs actual)",
        "4. Determine slab (tier lookup)",
        "5. KPI payout (base × slab rate)",
        "6. Apply weight (× weight %)",
        "<b>7. AGGREGATE KPIs + HELPER TRIP BONUS  ← integration point</b>",
        "8. Apply multipliers (growth, strategic)",
        "9. Apply penalties (high returns, audit fail)",
        "10. Apply caps (max per plan, % of salary)",
        "11. Store payout",
        "12. Create approval entry",
        "13. Complete run",
    ]
    for s in pipeline:
        story.append(Paragraph("• " + s, bullet_style))

    story.append(Paragraph("Helper integration details (Step 7):", body))
    for s in [
        "Pipeline queries trips joined with trip_participants for the employee and period",
        "Looks up the correct rate from helper_trip_rates using team size",
        "Multiplies rate × days_count per trip",
        "Sums across all trips",
        "Adds to gross_payout BEFORE multipliers/penalties/caps are applied",
        "Persists helper_trip_bonus + full breakdown in calculation_details JSON",
    ]:
        story.append(Paragraph("• " + s, bullet_style))

    story.append(Paragraph("Phase 4 — Approval Workflow", h2))
    story.append(Paragraph("submitted → manager_approved → finance_approved → hr_approved → locked", code_style))
    for s in [
        "Each stage requires a named approver",
        "Rejections require a mandatory reason",
        "Locked payouts cannot be edited",
        "Reopening requires super-admin privileges",
        "Every transition is logged in approval_log (immutable)",
    ]:
        story.append(Paragraph("• " + s, bullet_style))

    story.append(Paragraph("Phase 5 — Audit &amp; Governance", h2))
    story.append(mk_table([
        ["Entity Type", "Tracked Actions"],
        ["trip", "created, updated, deleted"],
        ["helper_rates", "updated (rate changes)"],
        ["employee_payouts", "calculated, approved, locked"],
    ], col_widths=[4 * cm, 10 * cm]))

    story.append(PageBreak())

    # 7. RATE LOOKUP
    story.append(Paragraph("7. Rate Lookup Logic", h1))
    story.append(Paragraph(
        "Tiered lookup — finds the highest team_size tier ≤ actual team size.", body))
    story.append(Paragraph("Behavior:", body))
    for s in [
        "Team of 1 → picks tier 1 (12 AED)",
        "Team of 2 → picks tier 2 (7 AED)",
        "Team of 3 → picks tier 3 (5 AED)",
        "Team of 5 → picks tier 4 (4 AED, since no tier 5 defined)",
        "Team of 100 → still picks tier 4 (highest defined ≤ 100)",
    ]:
        story.append(Paragraph("• " + s, bullet_style))
    story.append(Paragraph(
        "Plan-scoped rates override global defaults. If a plan has its own rates, "
        "those are used; otherwise the global defaults apply.", body))

    # 8. API REFERENCE
    story.append(Paragraph("8. API Reference", h1))
    story.append(mk_table([
        ["Method", "Endpoint", "Purpose"],
        ["GET", "/api/trips", "List trips (filters: employee_id, period, status)"],
        ["GET", "/api/trips/:id", "Single trip with participants"],
        ["POST", "/api/trips", "Create trip with participants"],
        ["PUT", "/api/trips/:id", "Update trip or participants"],
        ["DELETE", "/api/trips/:id", "Remove trip"],
        ["GET", "/api/trips/rates/:planId", "Get rate config"],
        ["PUT", "/api/trips/rates/:planId", "Update rate tiers"],
        ["GET", "/api/trips/commission/preview", "Preview commission without posting"],
    ], col_widths=[2 * cm, 5 * cm, 8.5 * cm]))

    # 9. EDGE CASES
    story.append(Paragraph("9. Edge Cases Handled", h1))
    story.append(mk_table([
        ["Scenario", "Behavior"],
        ["Trip with no end date", "Treated as 1-day (days_count = 1)"],
        ["End date before start date", "Falls back to 1 day"],
        ["Team size exceeds all tiers", "Uses the highest defined tier"],
        ["No rates configured for plan", "Falls back to global defaults"],
        ["Trip status = cancelled", "Excluded from commission calculation"],
        ["Employee on multiple trips same day", "All counted independently"],
        ["Rate changed mid-period", "Recalc uses current rate; locked plans preserve historical"],
        ["Empty trip log for period", "Returns total_commission: 0 cleanly"],
    ], col_widths=[5.5 * cm, 10 * cm]))

    # 10. INTEGRATION
    story.append(Paragraph("10. Integration with Existing Commission Logic", h1))
    story.append(Paragraph("The helper bonus flows through the standard pipeline:", body))
    story.append(mk_table([
        ["Pipeline Stage", "Effect on Helper Bonus"],
        ["Eligibility Check", "If employee fails min_sales, net = 0"],
        ["Multipliers", "Apply to gross (now includes helper)"],
        ["Penalties", "Deducts from total including helper"],
        ["Caps", "max_per_plan and percent_of_salary apply to final total"],
        ["Splits", "Role split applies to final total"],
        ["Approval", "Part of the same approval flow"],
        ["Audit trail", "Full JSON breakdown in calculation_details"],
    ], col_widths=[4 * cm, 11.5 * cm]))
    story.append(Paragraph(
        "<i>Important: If an employee's total is capped (e.g. 150% of salary), the helper bonus is added "
        "to gross before the cap is enforced. The cap may consume the helper bonus entirely — "
        "this matches how all other commission components behave.</i>", note_style))

    story.append(PageBreak())

    # 11. WORKED EXAMPLE
    story.append(Paragraph("11. Complete Worked Example", h1))
    story.append(Paragraph("<b>Setup:</b> Plan plan-01 'Salesman Monthly Incentive'", body))
    story.append(Paragraph("<b>Period:</b> 2026-04", body))
    story.append(Paragraph("<b>Employee:</b> Khalid Omar (emp-003), Van Sales Driver, salary 4,500 AED", body))

    story.append(Paragraph("Step 1 — Rates configured", h2))
    story.append(mk_table([
        ["Team Size", "Rate AED/day"],
        ["1", "12"], ["2", "7"], ["3", "5"], ["4+", "4"],
    ], col_widths=[4 * cm, 4 * cm]))

    story.append(Paragraph("Step 2 — Trips logged for April", h2))
    story.append(mk_table([
        ["Trip #", "Start", "End", "Days", "Team", "Rate", "Earned"],
        ["TRIP-APR-001", "Apr 5", "—", "1", "1 solo", "12", "12"],
        ["TRIP-APR-002", "Apr 6", "Apr 8", "3", "2 pair", "7", "21"],
        ["TRIP-APR-003", "Apr 9", "—", "1", "1 solo", "12", "12"],
        ["TEST-MULTIDAY", "Apr 10", "Apr 12", "3", "1 solo", "12", "36"],
        ["TOTAL", "", "", "8", "", "", "81"],
    ], col_widths=[3 * cm, 1.9 * cm, 1.9 * cm, 1.5 * cm, 2 * cm, 1.5 * cm, 2.2 * cm]))

    story.append(Paragraph("Step 3 — Calculation runs", h2))
    for s in [
        "Step 1: Fetch 104 transactions + 4 trips = 108 records",
        "Step 2: Apply filters, 104 pass",
        "Step 3: KPI Revenue: 82,341 / 50,000 = 164.68%",
        "Step 4: Tier 5 (120%+) → 12% rate",
        "Step 5: 2000 × 12% = 240",
        "Step 6: 240 × 40% weight = 96",
        "<b>Step 7: 48,636 KPI sum + 81 helper trip bonus = 48,717.60 gross</b>",
        "Step 8–10: No multipliers / penalties / caps hit",
        "Step 11–13: Stored, approval submitted, run completed",
    ]:
        story.append(Paragraph("• " + s, bullet_style))

    story.append(Paragraph("Step 4 — UI Result", h2))
    story.append(mk_table([
        ["Metric", "Value"],
        ["Gross Payout", "AED 48,717.60"],
        ["Helper Trip Bonus", "+AED 81 (4 trips · 8 days)"],
        ["Multiplier", "+AED 0"],
        ["Penalty", "-AED 0"],
        ["Net Payout", "AED 48,717.60"],
    ], col_widths=[5 * cm, 8 * cm]))

    # 12. ACCEPTANCE
    story.append(Paragraph("12. Acceptance Criteria (all met)", h1))
    for s in [
        "Rate table editable with visual live preview",
        "Solo rate is higher than pair rate (default: 12 vs 7)",
        "Team rates progressively lower (5 → 4)",
        "Multi-day trips multiply correctly (rate × days)",
        "Live form shows calculation before saving a trip",
        "Trip log filtered by period with 'all periods' option",
        "Pipeline Step 7 integrates helper bonus without breaking other steps",
        "Net payout visible in UI with helper bonus highlighted",
        "Approval workflow intact",
        "Audit trail captures helper commission details",
        "Works across multiple plans with plan-scoped overrides",
        "Commission Preview API for pre-calc verification",
    ]:
        story.append(Paragraph("✓ " + s, bullet_style))

    # 13. NFR
    story.append(Paragraph("13. Non-Functional Requirements", h1))
    story.append(mk_table([
        ["NFR", "How"],
        ["Scalability", "Trip queries indexed on employee_id + period"],
        ["Performance", "Rate lookup cached per run"],
        ["Auditability", "Full JSON breakdown persisted"],
        ["Configurability", "All rates editable via UI + API"],
        ["Multi-currency", "Rate table has currency column"],
        ["Multi-plan", "Plan-scoped overrides supported"],
        ["Reversibility", "Trip deletion supported"],
    ], col_widths=[4 * cm, 11.5 * cm]))

    # 14. FUTURE
    story.append(Paragraph("14. Future Enhancements (not in v1)", h1))
    for s in [
        "GPS validation — auto-reject trips without GPS confirmation",
        "Automatic trip creation — pull from telematics via webhook",
        "Trip-level approval before commission counts",
        "Rate variations by route/distance",
        "Seasonal rates (Ramadan, end-of-year)",
        "Bonus for perfect attendance",
    ]:
        story.append(Paragraph("• " + s, bullet_style))

    # 15. CONTACT
    story.append(Paragraph("15. Contact &amp; Environment", h1))
    story.append(mk_table([
        ["Field", "Value"],
        ["Project", "WINIT CommissionIQ Helper Trip Module"],
        ["Environment", "Production (Vercel, Mumbai region)"],
        ["Database", "Turso libSQL (ap-south-1)"],
        ["Deployment URL", "https://commission-app-six.vercel.app"],
        ["Status", "Live and operational"],
    ], col_widths=[4 * cm, 11.5 * cm]))

    story.append(Spacer(1, 12))
    story.append(Paragraph(
        "<para alignment='center'><i>— End of Document —</i></para>",
        ParagraphStyle("End", parent=body, textColor=colors.HexColor("#5B6B8C"))))

    doc.build(story)
    print(f"[OK] PDF written: {PDF_PATH}")


if __name__ == "__main__":
    build_docx()
    build_pdf()
